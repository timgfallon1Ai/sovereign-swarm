"""Word document extraction -- uses python-docx when available, stubs otherwise."""

from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from sovereign_swarm.document_intel.models import ExtractedContent

logger = structlog.get_logger()

try:
    from docx import Document  # type: ignore[import-untyped]

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


class DocxExtractor:
    """Extract text, tables, and metadata from Word documents."""

    def extract_text(self, path: str | Path) -> str:
        """Return the full text content of a DOCX file."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        if _HAS_DOCX:
            return self._text_docx(path)

        logger.warning("docx_extractor.no_library", path=str(path))
        return f"[STUB] DOCX text extraction not available -- install python-docx. Path: {path}"

    def extract_tables(self, path: str | Path) -> list[list[list[str]]]:
        """Return tables found in the document."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        if _HAS_DOCX:
            return self._tables_docx(path)

        return []

    def extract_metadata(self, path: str | Path) -> dict[str, Any]:
        """Return document metadata."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        if _HAS_DOCX:
            return self._metadata_docx(path)

        return {"stub": True, "path": str(path)}

    def extract_all(self, path: str | Path) -> ExtractedContent:
        """Full extraction -- text, tables, metadata."""
        path = Path(path)
        text = self.extract_text(path)
        tables = self.extract_tables(path)
        metadata = self.extract_metadata(path)
        return ExtractedContent(
            text=text,
            tables=tables,
            metadata=metadata,
            page_count=metadata.get("page_count", 0),
        )

    # ------------------------------------------------------------------
    # python-docx backends
    # ------------------------------------------------------------------

    @staticmethod
    def _text_docx(path: Path) -> str:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def _tables_docx(path: Path) -> list[list[list[str]]]:
        doc = Document(str(path))
        all_tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            all_tables.append(rows)
        return all_tables

    @staticmethod
    def _metadata_docx(path: Path) -> dict[str, Any]:
        doc = Document(str(path))
        props = doc.core_properties
        meta: dict[str, Any] = {
            "author": props.author or "",
            "title": props.title or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
        }
        meta["page_count"] = len(doc.paragraphs)  # approximate
        return meta
